import os
import imghdr
import logging
import requests
import validators
import urllib.request
import pypdfium2 as pdfium
from io import BytesIO
from PIL import Image
from docx import Document
from docx.document import Document as docs
from typing import Optional, List, Dict, Tuple
from requests import Response
from loggingConfig import setupLogging
from http.client import HTTPException
from pytesseract import image_to_string
from langdetect import detect, LangDetectException
import multiprocessing

setupLogging()
logger = logging.getLogger(__name__)

# Set the number of workers to use (adjust this based on your system's capabilities)
NUM_WORKERS = 6  # You can change this value to suit your needs

def renderPage(filePath: str, pageIndex: int) -> Optional[bytes]:
    """
    Renders a specific page of a PDF file as a JPEG image.
    Args:
        filePath (str): The path to the PDF file.
        pageIndex (int): The index of the page to render.
    Returns:
        Optional[bytes]: The rendered page as a JPEG image in bytes format, or None if an error occurred.
    """
    try:
        pdfFile = pdfium.PdfDocument(filePath)
        renderer = pdfFile.render(
            pdfium.PdfBitmap.to_pil,
            page_indices=[pageIndex],
            scale=(300/72)
        )
        imageList: List[Tuple] = list(renderer)
        image: Tuple = imageList[0]
        imageBytearray = BytesIO()
        image.save(imageBytearray, format='jpeg', optimize=True)
        return imageBytearray.getvalue()
    except Exception as e:
        logger.error("An error occurred while rendering page: %s", e, exc_info=True)
        return None

def convertPdfToImages(filePath: str) -> Optional[List[Dict[int, bytes]]]:
    """
    Converts a PDF file to a list of images.
    Args:
        filePath (str): The path to the PDF file.
    Returns:
        Optional[List[Dict[int, bytes]]]: A list of images represented as dictionaries,
        where the key is the page index and the value is the image data in bytes. Returns None
        if an error occurs during the conversion.
    """
    try:
        if imghdr.what(filePath) is not None:
            with open(filePath, 'rb') as f:
                return [{0: f.read()}]
        pdfFile = pdfium.PdfDocument(filePath)
        finalImages = []
        for i in range(len(pdfFile)):
            image = renderPage(filePath, i)
            if image:
                finalImages.append({i: image})
        return finalImages
    except Exception as e:
        logger.error("An error occurred while converting PDF to images: %s", e, exc_info=True)
        return None

def process_single_image(image_bytes: bytes) -> Optional[List[str]]:
    """
    Process a single image and extract text from it word by word.
    """
    try:
        image = Image.open(BytesIO(image_bytes))
        rawText = image_to_string(image, lang='eng+hin+te')
        
        if not rawText or len(rawText) < 10:
            return None
        
        try:
            language = detect(rawText)
            if language != 'en':
                return None
        except LangDetectException as e:
            logger.error("An error occurred while detecting language: %s", e, exc_info=True)
            return None
        
        word_list = rawText.split()
        return word_list if word_list else None
    except Exception as e:
        logger.error("An error occurred while processing single image: %s", e, exc_info=True)
        return None

def processImage(image_bytes_list: List[bytes]) -> Optional[List[str]]:
    """
    Process the given images and extract text from them word by word using parallel processing.
    Args:
        image_bytes_list (List[bytes]): A list of image bytes.
    Returns:
        Optional[List[str]]: A list of extracted words from all images, or None if the text could not be extracted.
    """
    try:
        # Use multiprocessing to process images in parallel with a limited number of workers
        with multiprocessing.Pool(processes=NUM_WORKERS) as pool:
            results = pool.map(process_single_image, image_bytes_list)
        
        # Combine results from all images
        all_words = [word for result in results if result for word in result]
        return all_words if all_words else None
    except Exception as e:
        logger.error("An error occurred while processing images: %s", e, exc_info=True)
        return None

def extractTextFromPdf(filePath: str) -> Optional[str]:
    """
    Extracts text from a PDF file given its file path.
    Args:
        filePath (str): The path to the PDF file.
    Returns:
        Optional[str]: The extracted text from the PDF file, or None if an error occurred.
    """
    try:
        images: Optional[List[Dict[int, bytes]]] = convertPdfToImages(filePath)
        image_bytes_list = []
        if images:
            for image_dict in images:
                for _, image_bytes in image_dict.items():
                    image_bytes_list.append(image_bytes)
        
        extractedData = processImage(image_bytes_list)
        if extractedData:
            return ' '.join(extractedData)
        return None
    except Exception as e:
        logger.error("An error occurred while extracting text from PDF file: %s", e, exc_info=True)
        return None

def extractTextFromPdfUrl(url: str) -> Optional[str]:
    """
    Extracts text from a PDF file given its URL.
    Args:
        url (str): The URL of the PDF file.
    Returns:
        Optional[str]: The extracted text from the PDF file, or None if an error occurred.
    """
    try:
        if not validators.url(url):
            raise HTTPException("Invalid URL")
        temp_file_path, _ = urllib.request.urlretrieve(url)
        return extractTextFromPdf(temp_file_path)
    except (ValueError, HTTPException) as e:
        logger.error("An error occurred while extracting text from PDF URL: %s", e, exc_info=True)
        return None
    except Exception as e:
        logger.error("An error occurred while extracting text from PDF URL: %s", e, exc_info=True)
        return None

def extractTextFromDocxUrl(docxUrl: str) -> Optional[str]:
    """
    Extracts text from a DOCX file given its URL.
    Args:
        docxUrl (str): The URL of the DOCX file.
    Returns:
        Optional[str]: The extracted text from the DOCX file, or None if an error occurred.
    """
    try:
        response: Response = requests.get(docxUrl)
        docxContent = response.content
        docxFile = BytesIO(docxContent)
        doc: docs = Document(docxFile)
        extractedText: str = ""
        for paragraph in doc.paragraphs:
            extractedText += paragraph.text
        return extractedText
    except Exception as e:
        logger.error("An error occurred while extracting text from DOCX URL: %s", e, exc_info=True)
        return None

def extractTextFromDocx(filePath: str) -> Optional[str]:
    """
    Extracts text from a DOCX file given its file path.
    Args:
        filePath (str): The path to the DOCX file.
    Returns:
        Optional[str]: The extracted text from the DOCX file, or None if an error occurred.
    """
    try:
        doc: docs = Document(filePath)
        extractedText: str = ""
        for paragraph in doc.paragraphs:
            extractedText += paragraph.text
        return extractedText
    except Exception as e:
        logger.error("An error occurred while extracting text from DOCX file: %s", e, exc_info=True)
        return None

def check_url_response(url: str) -> Optional[str]:
    """
    Check if the URL is accessible and return the final URL after following redirects.
    
    Args:
        url (str): The URL to check.
    
    Returns:
        Optional[str]: The final URL if accessible, None otherwise.
    """
    try:
        response = requests.head(url, allow_redirects=True, timeout=10)
        if response.status_code == 200:
            return response.url
        else:
            logger.warning(f"URL {url} returned status code {response.status_code}")
            return None
    except requests.RequestException as e:
        logger.error(f"Error checking URL {url}: {str(e)}")
        return None

def extractContent(url_or_path: str) -> Optional[str]:
    """
    Extracts text content from a file given its URL or file path.
    Args:
        url_or_path (str): The URL or file path of the file.
    Returns:
        Optional[str]: The extracted text content from the file, or None if an error occurred.
    """
    try:
        if validators.url(url_or_path):
            final_url = check_url_response(url_or_path)
            if not final_url:
                logger.error(f"Unable to access URL: {url_or_path}")
                return None
            url_or_path = final_url

        if url_or_path.endswith('.pdf'):
            if validators.url(url_or_path):
                return extractTextFromPdfUrl(url_or_path)
            elif os.path.isfile(url_or_path):
                return extractTextFromPdf(url_or_path)
        elif url_or_path.endswith('.docx'):
            if validators.url(url_or_path):
                return extractTextFromDocxUrl(url_or_path)
            elif os.path.isfile(url_or_path):
                return extractTextFromDocx(url_or_path)
        # elif url_or_path.endswith('.doc'):
        #     return extractTextFromDocUrl(url_or_path)
        # Uncomment and complete the following lines if needed
        # elif url_or_path.endswith('.xls') or url_or_path.endswith('.xlsx'):
        #     return convertFromExcel(url_or_path)
        else:
            logger.error("Invalid file format")
            return None
    except Exception as e:
        logger.error(f"An error occurred while extracting content: {e}", exc_info=True)
        return None

# Uncomment and complete the following function if needed
# def convertFromExcel(url: str) -> Optional[List[Dict]]:
#     """
#     Converts an Excel file to a list of dictionaries.
# 
#     Args:
#         url (str): The URL of the Excel file.
# 
#     Returns:
#         Optional[List[Dict]]: The data from the Excel file as a list of dictionaries, or None if an error occurred.
#     """
#     try:
#         response = requests.get(url)
#         excel_file = BytesIO(response.content)
#         reader = pd.read_excel(excel_file, header=0)
#         reader = reader.dropna(axis=1, how='all')
#         header_row = reader.iloc[0]
#         reader = reader[1:]
#         header_row = header_row.astype(str)
#         reader.columns = header_row
#         excel_data = reader.to_dict(orient='records')
#         return excel_data
#     except Exception as e:
#         logger.error("An error occurred while converting from Excel: %s", e, exc_info=True)
#         return None

# def extractTextFromDocUrl(url: str) -> Optional[str]:
#     """
#     Extracts text from a DOC file given its URL.
#
#     Args:
#         url (str): The URL of the DOC file.
#
#     Returns:
#         Optional[str]: The extracted text from the DOC file, or None if an error occurred.
#     """
#     try:
#         temp_file_path, _ = urllib.request.urlretrieve(url)
#         extractedText: str = textract.process(temp_file_path, output_encoding='utf_8', extension='doc').decode('utf-8')
#         return extractedText
#     except Exception as e:
#         logger.error("An error occurred while extracting text from DOC URL: %s", e, exc_info=True)
#         return None